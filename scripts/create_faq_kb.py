#!/usr/bin/env python3
"""
Generate a .docx KB document from https://www.novaclinic.ca/faq
Covers: insurance, direct billing, Alberta Health coverage, cancellation policy.
"""
from docx import Document
from pathlib import Path

doc = Document()

# ── Insurance & Billing ───────────────────────────────────────────────
doc.add_heading("INSURANCE & BILLING — FAQ", level=1)

doc.add_heading("Do You Offer Direct Billing?", level=2)
doc.add_paragraph(
    "Naturopathic health services are covered by most major extended healthcare "
    "insurance plans. For insurance companies that allow direct billing, Nova Clinic "
    "handles the billing directly. Otherwise, patients pay upfront and submit "
    "claims to their insurance company — the clinic will assist with any paperwork needed."
)

doc.add_heading("Direct Billing Insurance Companies", level=2)
doc.add_paragraph(
    "Nova Clinic currently offers direct billing to the following insurance companies: "
    "Manulife, Blue Cross, Desjardins, and Provider Connect. "
    "If your insurance company is not on this list, you can pay upfront and "
    "submit your receipt for reimbursement."
)

doc.add_heading("Will My Treatment Be Fully Covered by Insurance?", level=2)
doc.add_paragraph(
    "Coverage percentages vary by insurer and individual policy. Some insurance "
    "companies cover only certain percentages or impose daily maximums. Patients "
    "are responsible for paying any outstanding amounts upfront. Please confirm "
    "with your insurance company before your visit regarding your specific "
    "coverage limits and what services are included."
)

doc.add_heading("Does Alberta Health Care Cover Naturopathic Services?", level=2)
doc.add_paragraph(
    "The Alberta Health Care Insurance Plan (AHCIP) does NOT currently cover "
    "naturopathic services in Alberta. However, most private insurance companies "
    "typically cover naturopathic services under their Extended Health Coverage "
    "plans, as long as the practitioner is a member of the College of Naturopathic "
    "Doctors of Alberta (CNDA)."
)

doc.add_heading("Direct Billing Disclaimer", level=2)
doc.add_paragraph(
    "The clinic provides a detailed breakdown of services rendered. It is the "
    "client's responsibility to ensure that their insurance company covers each "
    "of the naturopathic services rendered. If a claim is marked as "
    "'Pending/for Review' by the insurance company, the patient is responsible "
    "for payment, and the clinic will void the claim for re-submission. "
    "Patients must also pay when insurance denies a claim, coverage is maxed out, "
    "or a plan terminates."
)

# ── Cancellation Policy ──────────────────────────────────────────────
doc.add_heading("CANCELLATION & RESCHEDULING POLICY — FAQ", level=1)

doc.add_heading("24-Hour Cancellation Policy", level=2)
doc.add_paragraph(
    "Nova Clinic has a 24-hour cancellation policy. Please inform the clinic at "
    "least 24 hours before your scheduled appointment if you need to cancel or "
    "reschedule. Failure to properly inform the clinic of any appointment "
    "cancellations within this period may result in a fee being added to your account."
)
doc.add_paragraph(
    "To reschedule or cancel your appointment, contact the clinic at "
    "(587) 391-5753, email admin@novaclinic.ca, or manage your appointment "
    "through the online booking portal at novacliniccalgary.janeapp.com."
)

# ── Save ──────────────────────────────────────────────────────────────
out_path = Path("kb/sources/nova_clinic_faq.docx")
doc.save(str(out_path))
print(f"Created {out_path} successfully!")
