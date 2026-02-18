#!/usr/bin/env python3
"""
Generate a supplementary .docx KB document with pricing and service info
that was missing from the main KB but is listed on the Nova Clinic fees page.
"""
from docx import Document
from pathlib import Path

doc = Document()

# ── Osteopathic Manual Therapy ───────────────────────────────────────
doc.add_heading("SERVICE: OSTEOPATHIC MANUAL THERAPY", level=1)

doc.add_heading("Osteopathic Manual Therapy — Service Overview", level=2)
doc.add_paragraph(
    "Osteopathic Manual Therapy at Nova Clinic is a hands-on treatment approach that "
    "focuses on the musculoskeletal system to improve overall health and well-being. "
    "Osteopathic practitioners use gentle manual techniques to address structural "
    "imbalances, improve mobility, and support the body's natural healing processes."
)

doc.add_heading("Osteopathic Manual Therapy — Duration & Pricing", level=2)
doc.add_paragraph(
    "All pricing is sourced from the official Fees page and is subject to change."
)
doc.add_paragraph(
    "Initial Osteopathic Assessment: $150 (approximately 50 minutes). "
    "This first visit includes a thorough assessment of the patient's condition, "
    "medical history, and treatment planning."
)
doc.add_paragraph(
    "Osteopathic Follow-Up Treatment: $95 (approximately 25 minutes). "
    "Follow-up visits focus on ongoing treatment and progress evaluation."
)
doc.add_paragraph(
    "Children (Age 12 & Under) — Initial Assessment: $90. "
    "Children (Age 12 & Under) — Follow-Up Treatment: $75."
)

# ── Kids Massage Pricing ────────────────────────────────────────────
doc.add_heading("Kids Massage Therapy (Age 12 & Under) — Pricing", level=2)
doc.add_paragraph(
    "All pricing is sourced from the official Fees page and is subject to change. "
    "Kids massage sessions are specifically designed for children aged 12 and under."
)
doc.add_paragraph(
    "Kids Massage 30 minutes: $75. "
    "Kids Massage 45 minutes: $95. "
    "Kids Massage 60 minutes: $110."
)

# ── Prenatal/Postnatal Massage Pricing ──────────────────────────────
doc.add_heading("Prenatal & Postnatal Massage — Pricing", level=2)
doc.add_paragraph(
    "All pricing is sourced from the official Fees page and is subject to change. "
    "Prenatal and postnatal massage is available for patients who are 12 weeks "
    "or more into their pregnancy, or postpartum."
)
doc.add_paragraph(
    "Prenatal/Postnatal Massage 45 minutes: $100. "
    "Prenatal/Postnatal Massage 60 minutes: $120. "
    "Prenatal/Postnatal Massage 90 minutes: $160."
)

# ── Prolotherapy Pricing ────────────────────────────────────────────
doc.add_heading("Prolotherapy — Pricing", level=2)
doc.add_paragraph(
    "All pricing is sourced from the official Fees page and is subject to change. "
    "Prolotherapy sessions are approximately 15–30 minutes and cost $150–$200 per session. "
    "Prolotherapy is a regenerative injection therapy for joints and soft tissue. "
    "An Initial Injection Consultation (starting at $290, approximately 80 minutes) "
    "is required before prolotherapy treatment can begin."
)

# ── Acupuncture New Patient Add-On ──────────────────────────────────
doc.add_heading("Acupuncture — New Patient Add-On", level=2)
doc.add_paragraph(
    "New acupuncture patients can add an extra 15 minutes to their first Classic "
    "Acupuncture session for an additional $10. This allows extra time for intake "
    "assessment and discussion of health concerns. The standard Classic Acupuncture "
    "session is 45–60 minutes at $100–$120."
)

# ── Save ─────────────────────────────────────────────────────────────
out_path = Path("kb/sources/nova_clinic_supplementary_fees.docx")
doc.save(str(out_path))
print(f"Created {out_path} successfully!")
