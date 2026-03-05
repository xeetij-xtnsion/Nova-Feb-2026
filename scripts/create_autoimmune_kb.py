#!/usr/bin/env python3
"""
Generate a .docx KB document for Dr. Nurani's autoimmune specialty.
"""
from docx import Document
from pathlib import Path

doc = Document()

# ── Header ────────────────────────────────────────────────────────────
doc.add_heading("AUTOIMMUNE CONDITIONS — DR. ALI NURANI, ND", level=1)

doc.add_heading("Overview", level=2)
doc.add_paragraph(
    "Dr. Ali Nurani has a very strong focus in treating immune-related conditions, "
    "including all autoimmune conditions such as lupus. Autoimmune conditions require "
    "a comprehensive workup and plan by an experienced naturopathic doctor, such as "
    "Dr. Ali Nurani. At Nova Clinic, we take a thorough, root-cause approach to "
    "autoimmune disease — identifying triggers, reducing inflammation, supporting "
    "immune regulation, and creating personalized treatment plans."
)

doc.add_paragraph(
    "If you are dealing with an autoimmune condition or suspect you may have one, "
    "we recommend booking an Initial Naturopathic Consultation with Dr. Ali Nurani "
    "so he can conduct a comprehensive assessment. You can also start with a free "
    "15-minute Meet & Greet to discuss your concerns."
)

# ── Joint, Muscle, and Connective Tissue ─────────────────────────────
doc.add_heading("Joint, Muscle, and Connective Tissue Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Dr. Nurani treats a wide range of joint, muscle, and connective tissue "
    "autoimmune conditions, including: Rheumatoid Arthritis (RA), Psoriatic Arthritis, "
    "Ankylosing Spondylitis, Systemic Lupus Erythematosus (SLE), Sjögren's Syndrome, "
    "Scleroderma (Systemic Sclerosis), Dermatomyositis, Polymyositis, "
    "Mixed Connective Tissue Disease (MCTD), Polymyalgia Rheumatica, "
    "Relapsing Polychondritis, Palindromic Rheumatism, Adult-Onset Still's Disease, "
    "Reactive Arthritis (Reiter's Syndrome), Juvenile Idiopathic Arthritis (JIA), "
    "Undifferentiated Connective Tissue Disease (UCTD), Inclusion Body Myositis, "
    "Fibromyalgia (often co-occurs with immune-mediated components), "
    "Behçet's Disease, and CREST Syndrome (limited scleroderma)."
)

# ── Endocrine (Hormonal) ─────────────────────────────────────────────
doc.add_heading("Endocrine (Hormonal) Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Autoimmune conditions affecting the endocrine system that Dr. Nurani treats "
    "include: Hashimoto's Thyroiditis, Graves' Disease, Type 1 Diabetes Mellitus, "
    "Addison's Disease (adrenal insufficiency), Autoimmune Hypophysitis (pituitary gland), "
    "Autoimmune Pancreatitis, Autoimmune Polyglandular Syndrome Type 1 (APS-1) and "
    "Type 2 (APS-2), Autoimmune Oophoritis (premature ovarian failure), "
    "Autoimmune Orchitis (testicular inflammation), and Ord's Thyroiditis "
    "(variant of Hashimoto's)."
)

# ── Gastrointestinal and Liver ───────────────────────────────────────
doc.add_heading("Gastrointestinal and Liver Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Dr. Nurani treats gastrointestinal and liver autoimmune conditions including: "
    "Celiac Disease, Crohn's Disease, Ulcerative Colitis, Autoimmune Hepatitis "
    "(Type 1 & 2), Primary Biliary Cholangitis (PBC), Primary Sclerosing Cholangitis "
    "(PSC), Pernicious Anemia (atrophic gastritis), Microscopic Colitis, "
    "Eosinophilic Esophagitis (EoE), and Whipple's Disease."
)

# ── Neurological and Neuromuscular ───────────────────────────────────
doc.add_heading("Neurological and Neuromuscular Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Neurological and neuromuscular autoimmune conditions that Dr. Nurani can help "
    "with include: Multiple Sclerosis (MS), Myasthenia Gravis, Guillain-Barré "
    "Syndrome (GBS), Chronic Inflammatory Demyelinating Polyneuropathy (CIDP), "
    "Neuromyelitis Optica (NMO), Lambert-Eaton Myasthenic Syndrome (LEMS), "
    "Stiff Person Syndrome, Transverse Myelitis, Acute Disseminated "
    "Encephalomyelitis (ADEM), PANDAS, Autoimmune Encephalitis, Sydenham's Chorea, "
    "Narcolepsy Type 1, Paraneoplastic Neurological Syndromes, and Isaac's Syndrome "
    "(Neuromyotonia)."
)

# ── Dermatological (Skin) ────────────────────────────────────────────
doc.add_heading("Dermatological (Skin) Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Skin-related autoimmune conditions treated include: Psoriasis, Vitiligo, "
    "Alopecia Areata, Pemphigus Vulgaris, Bullous Pemphigoid, Cicatricial Pemphigoid, "
    "Dermatitis Herpetiformis (linked to Celiac), Erythema Nodosum, Lichen Planus, "
    "Lichen Sclerosus, Pyoderma Gangrenosum, and Epidermolysis Bullosa Acquisita."
)

# ── Hematological (Blood) ────────────────────────────────────────────
doc.add_heading("Hematological (Blood) Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Blood-related autoimmune conditions include: Immune Thrombocytopenic Purpura "
    "(ITP), Autoimmune Hemolytic Anemia, Antiphospholipid Syndrome (APS), "
    "Pure Red Cell Aplasia, Aplastic Anemia (some forms), Paroxysmal Nocturnal "
    "Hemoglobinuria (PNH), Evans Syndrome, and Autoimmune Neutropenia."
)

# ── Cardiovascular and Pulmonary ─────────────────────────────────────
doc.add_heading("Cardiovascular and Pulmonary Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Cardiovascular and pulmonary autoimmune conditions include: Vascular Vasculitis "
    "(Giant Cell Arteritis), Granulomatosis with Polyangiitis (GPA/Wegener's), "
    "Microscopic Polyangiitis, Eosinophilic Granulomatosis with Polyangiitis "
    "(Churg-Strauss), Kawasaki Disease, Takayasu's Arteritis, Polyarteritis Nodosa, "
    "Buerger's Disease, Goodpasture Syndrome (kidneys and lungs), Sarcoidosis, "
    "Idiopathic Pulmonary Fibrosis, and Dressler's Syndrome."
)

# ── Ocular and Auditory (Eyes and Ears) ──────────────────────────────
doc.add_heading("Ocular and Auditory Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Autoimmune conditions affecting the eyes and ears include: Uveitis, Scleritis, "
    "Cogan's Syndrome, Autoimmune Inner Ear Disease (AIED), Mooren's Ulcer, "
    "and Vogt-Koyanagi-Harada Syndrome."
)

# ── Renal and Other Systemic ─────────────────────────────────────────
doc.add_heading("Renal and Other Systemic Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Renal and systemic autoimmune conditions include: IgA Nephropathy (Berger's "
    "Disease), Membranous Nephropathy, Lupus Nephritis, Retroperitoneal Fibrosis, "
    "Systemic Juvenile Idiopathic Arthritis, and Amyloidosis (secondary to chronic "
    "inflammation)."
)

# ── Naturopathic Approach ────────────────────────────────────────────
doc.add_heading("Our Naturopathic Approach to Autoimmune Conditions", level=2)
doc.add_paragraph(
    "Dr. Nurani's approach to autoimmune conditions involves a comprehensive workup "
    "that may include advanced functional testing such as the Multiple Autoimmune "
    "Reactivity Screen (Antibody Array 5), food sensitivity testing, comprehensive "
    "gut analysis, and other specialized labs. Treatment plans are personalized and "
    "may include dietary modifications, targeted supplementation, IV nutrient therapy, "
    "lifestyle interventions, and stress management — all aimed at modulating the "
    "immune response, reducing inflammation, and addressing root causes."
)

# ── Cancer Co-Management ──────────────────────────────────────────────
doc.add_heading("Cancer Co-Management with Dr. Ali Nurani", level=2)
doc.add_paragraph(
    "While our naturopathic doctors do not focus on cancer, one of our naturopathic "
    "doctors, Dr. Ali Nurani, can co-manage some of the naturopathic treatments "
    "alongside conventional cancer care. This includes diet and nutritional guidance, "
    "various cancer-focused IV nutrition therapies, and botanical/herbal treatments. "
    "Dr. Nurani works alongside your oncology team to support your overall health "
    "and well-being during cancer treatment."
)
doc.add_paragraph(
    "If you or a loved one is going through cancer treatment and would like naturopathic "
    "support, we recommend booking an Initial Naturopathic Consultation with Dr. Ali "
    "Nurani to discuss how we can help complement your care plan."
)

# ── Save ──────────────────────────────────────────────────────────────
out_path = Path("kb/sources/nova_clinic_autoimmune.docx")
doc.save(str(out_path))
print(f"Created {out_path} successfully!")
