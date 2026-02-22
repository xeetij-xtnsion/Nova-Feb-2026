#!/usr/bin/env python3
"""
Generate a comprehensive .docx KB document with all pricing from
https://www.novaclinic.ca/fees — to be ingested into the knowledge base.

This replaces the older supplementary_fees doc with a single complete source.
"""
from docx import Document
from pathlib import Path

doc = Document()

# ── Naturopathic Services ─────────────────────────────────────────────
doc.add_heading("SERVICE: NATUROPATHIC MEDICINE — FEES & PRICING", level=1)

doc.add_heading("Initial Naturopathic Consultation — Pricing", level=2)
doc.add_paragraph(
    "The Initial Naturopathic Consultation is for first-time patients only. "
    "Duration: approximately 80 minutes. Price: $295. "
    "This visit includes a detailed case history, review of lab work, focused "
    "physical exam, review of current medications and supplements, and initiation "
    "of a personalized treatment plan."
)

doc.add_heading("Initial Injection/IV Consultation — Pricing", level=2)
doc.add_paragraph(
    "The Initial Injection/IV Consultation is for first-time patients interested "
    "in IV nutrient infusions, prolotherapy, or vitamin shots. "
    "Duration: approximately 80 minutes. Price: starting at $290. "
    "This includes assessment, safety review, and imaging/blood work review. "
    "The cost may or may not include the first treatment depending on your "
    "naturopathic doctor's individual plan."
)

doc.add_heading("Naturopathic Follow-Up Visit — Pricing", level=2)
doc.add_paragraph(
    "Follow-up visits for existing naturopathic patients. "
    "Duration: 15–45 minutes. Price: $90–$185. "
    "Includes assessment, safety review, imaging/bloodwork review, "
    "and treatment discussion."
)

doc.add_heading("Extended Naturopathic Follow-Up Visit — Pricing", level=2)
doc.add_paragraph(
    "For existing patients who need more time to discuss their health concerns. "
    "Duration: 60 minutes. Price: $225."
)

doc.add_heading("IV Nutrient Therapy — Pricing", level=2)
doc.add_paragraph(
    "IV Nutrient Therapy at Nova Clinic. Pricing effective January 1st, 2022. "
    "IV Drip: $85–$195 (45 minutes to 2 hours). "
    "IV Push: $65–$115 (15–45 minutes). "
    "Individualized infusions may include glutathione, NAD+, vitamin C, "
    "magnesium, B12, and more. An Initial Injection/IV Consultation is "
    "required before starting IV therapy."
)

doc.add_heading("Vitamin IM (Intramuscular) Injection — Pricing", level=2)
doc.add_paragraph(
    "Quick vitamin injection delivered into the gluteal or deltoid muscle. "
    "Duration: less than 10 minutes. Price: $40–$75. "
    "No IV access is required for intramuscular injections."
)

doc.add_heading("Prolotherapy — Pricing", level=2)
doc.add_paragraph(
    "Prolotherapy is a regenerative injection therapy using a natural irritant "
    "(dextrose) injected into soft tissue of injured joints, ligaments, or tendons "
    "to stimulate healing. "
    "Duration: 15–30 minutes. Price: $150–$200 per session. "
    "An Initial Injection Consultation (starting at $290, approximately 80 minutes) "
    "is required before prolotherapy treatment can begin."
)

doc.add_heading("Trigger Point Injection — Pricing", level=2)
doc.add_paragraph(
    "Trigger Point Injections treat painful muscle areas containing trigger points. "
    "Used for chronic musculoskeletal disorders. "
    "Duration: 15–30 minutes. Price: $150–$200 per session."
)

doc.add_heading("Functional Testing — Pricing", level=2)
doc.add_paragraph(
    "Functional testing prices vary depending on the type of test. "
    "Available tests include standard diagnostic panels and specialized testing "
    "such as hormone testing, stool analysis, food sensitivity panels, "
    "nutrient testing, urine minerals, and heavy metal testing. "
    "Please consult with your naturopathic doctor for specific pricing."
)

# ── Osteopathic Services ──────────────────────────────────────────────
doc.add_heading("SERVICE: OSTEOPATHIC MANUAL THERAPY — FEES & PRICING", level=1)

doc.add_heading("Osteopathic Manual Therapy — Pricing", level=2)
doc.add_paragraph(
    "Initial Osteopathic Assessment: $150 (approximately 50 minutes). "
    "Includes health history intake, assessment, followed by treatment. "
    "Osteopathic treatment removes blockages and balances restrictions."
)
doc.add_paragraph(
    "Osteopathic Follow-Up Treatment: $95 (approximately 25 minutes)."
)
doc.add_paragraph(
    "Children (Age 12 & Under) — Initial Assessment: $90 (approximately 50 minutes). "
    "Children (Age 12 & Under) — Follow-Up: $75 (approximately 25 minutes)."
)

# ── Massage Therapy ───────────────────────────────────────────────────
doc.add_heading("SERVICE: MASSAGE THERAPY — FEES & PRICING", level=1)

doc.add_heading("Custom Massage — Pricing", level=2)
doc.add_paragraph(
    "Custom massage sessions can be tailored to relaxation, deep tissue, "
    "therapeutic, lymphatic drainage, and more. "
    "30 minutes: $75. "
    "45 minutes: $100. "
    "60 minutes: $120. "
    "75 minutes: $140. "
    "90 minutes: $160."
)

doc.add_heading("Kids Massage (Age 12 & Under) — Pricing", level=2)
doc.add_paragraph(
    "Massage sessions designed for children aged 12 and under. "
    "Parental consent is required before treatment. Parents are welcome in the room. "
    "30 minutes: $75. "
    "45 minutes: $95. "
    "60 minutes: $110."
)

doc.add_heading("Prenatal & Postnatal Massage — Pricing", level=2)
doc.add_paragraph(
    "Available for patients 12 weeks or more into pregnancy, or postpartum. "
    "45 minutes: $100. "
    "60 minutes: $120. "
    "90 minutes: $160."
)

doc.add_heading("Massage Add-Ons — Pricing", level=2)
doc.add_paragraph(
    "Add-on services available with any massage booking: "
    "Hot Stone: $35. "
    "Hydrotherapy: $35. "
    "Suction Cupping: $50."
)

# ── Acupuncture / TCM ────────────────────────────────────────────────
doc.add_heading("SERVICE: ACUPUNCTURE & TCM — FEES & PRICING", level=1)

doc.add_heading("Classic Acupuncture — Pricing", level=2)
doc.add_paragraph(
    "Classic Acupuncture is a form of Traditional Chinese Medicine (TCM) used to "
    "naturally improve the flow of Qi (energy) throughout the body. "
    "Duration: 45–60 minutes. Price: $100–$120. "
    "The initial treatment includes a detailed diagnosis utilizing many philosophies "
    "of Traditional Chinese Medicine. "
    "Applications include digestive issues, menstrual cycle support, fertility, "
    "pain management, stress/anxiety, sleep issues, and headaches."
)
doc.add_paragraph(
    "New Patient Add-On: an additional $10 for an extra 15 minutes on your first visit, "
    "allowing more time for intake assessment and discussion of health concerns."
)
doc.add_paragraph(
    "Free discovery call consultations are available for acupuncture."
)

# ── Cancellation & Rescheduling Policy ────────────────────────────────
doc.add_heading("CANCELLATION & RESCHEDULING POLICY", level=1)

doc.add_heading("24-Hour Cancellation Policy", level=2)
doc.add_paragraph(
    "Nova Clinic has a 24-hour cancellation policy. Please inform the clinic at "
    "least 24 hours before your scheduled appointment if you need to cancel or "
    "reschedule. Failure to properly inform the clinic of any appointment "
    "cancellations within this period may result in a fee being added to your account. "
    "This policy applies to all services including naturopathic consultations, "
    "osteopathic manual therapy, massage therapy, and acupuncture."
)
doc.add_paragraph(
    "To reschedule or cancel your appointment, contact the clinic at "
    "(587) 391-5753, email admin@novaclinic.ca, or manage your appointment "
    "through the online booking portal at novacliniccalgary.janeapp.com."
)

# ── Save ──────────────────────────────────────────────────────────────
out_path = Path("kb/sources/nova_clinic_fees.docx")
doc.save(str(out_path))
print(f"Created {out_path} successfully!")
