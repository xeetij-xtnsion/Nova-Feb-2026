#!/usr/bin/env python3
"""
Generate a .docx knowledge base document for Nova Clinic's functional testing services.
Data sourced from https://www.novaclinic.ca functional testing pages (Feb 2026).
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path

doc = Document()

# ── Overview ──────────────────────────────────────────────────────────
doc.add_heading("Testing at Nova Clinic — Overview", level=1)
doc.add_paragraph(
    "Testing ordered by the naturopathic doctor is not subsidized by the government of "
    "Alberta and will be an out-of-pocket expense, but can be billed to extended health "
    "insurance. Both types of testing require a naturopathic doctor consultation and will "
    "be ordered only if the doctor feels that it is recommended and necessary."
)

doc.add_heading("Two Categories of Testing", level=2)
doc.add_paragraph(
    "There are two major categories of testing at Nova Clinic:\n\n"
    "1. Standardized Testing — This is similar to what your family doctor is able to "
    "order through Alberta Health Services. These are conventional lab tests such as "
    "blood work, thyroid panels, cholesterol, blood sugar, and other routine diagnostics.\n\n"
    "2. Functional Testing — These tests measure 'sub-clinical' imbalances — areas where "
    "the body isn't 'broken' yet but is struggling to maintain balance. Functional tests "
    "use blood, stool, saliva, or dried urine to capture data and provide extremely deep "
    "insights about one's health. Nova Clinic has partnered with the best functional labs "
    "in the world to be able to offer the very best functional tests for our patients."
)

doc.add_heading("Functional Testing — Detailed Test Catalog", level=2)
doc.add_paragraph(
    "Nova Clinic offers 12 specialized functional and diagnostic tests to help patients "
    "and practitioners gain deeper insight into specific areas of health. All functional "
    "tests require an initial naturopathic consultation before testing can be ordered. "
    "Tests are not available as standalone bookings. Results are reviewed during a "
    "follow-up naturopathic appointment."
)

# ── 1. Food Sensitivity Testing ──────────────────────────────────────
doc.add_heading("IgG Food Sensitivity Test", level=2)
doc.add_paragraph(
    "The IgG Food Sensitivity Test is a blood test that measures immune system reactions "
    "to foods by detecting elevated IgG antibodies. It helps identify foods that may be "
    "triggering symptoms like digestive issues, migraines, inflammation, and weight gain. "
    "IgG reactions are dose-dependent — foods most frequently consumed or consumed in "
    "large quantity are more likely to produce symptoms."
)
doc.add_paragraph(
    "This test is ideal for patients experiencing digestive disorders, unexplained weight "
    "gain, migraines, chronic inflammation, or recurring health issues potentially linked "
    "to food reactions."
)
doc.add_paragraph(
    "The test is a simple finger-prick blood draw performed in-office. A small puncture "
    "is made on the finger and a few drops of blood are collected into a microtainer "
    "vessel and sent to the lab. No fasting is required. The procedure takes approximately "
    "10 minutes."
)
doc.add_paragraph(
    "Test results are usually available 7–10 days after collection. The half-life of IgG "
    "antibodies is 23–96 days, so complete clearance of reactive foods requires 3–18 "
    "months after elimination."
)
doc.add_paragraph(
    "Pricing for the IgG Food Sensitivity Test: 120+ food antigens panel costs $275, "
    "200+ food antigens panel costs $379, 160+ vegetarian food antigens panel costs $315. "
    "Add-on antigens are available at $20 each. A naturopathic consultation is required "
    "before testing. Insurance coverage varies — doctor fees may be covered under "
    "extended health plans."
)

# ── 2. DUTCH Hormone Test ────────────────────────────────────────────
doc.add_heading("DUTCH Hormone Test", level=2)
doc.add_paragraph(
    "The DUTCH Test (Dried Urine Test for Comprehensive Hormones) is a non-invasive, "
    "comprehensive method for assessing hormone levels through dried urine samples "
    "collected over a 24-hour period. It provides a more complete picture of hormone "
    "balance compared to traditional blood tests that only offer a single-point snapshot."
)
doc.add_paragraph(
    "The DUTCH Complete test measures 18 separate markers including sex hormones "
    "(estrogen, progesterone, testosterone), adrenal hormones (DHEA, cortisol), "
    "neurotransmitter metabolites (serotonin, dopamine, melatonin, norepinephrine), "
    "and oxidative stress markers. Note: thyroid hormones are NOT tested by the DUTCH "
    "test and are better assessed via blood tests."
)
doc.add_paragraph(
    "This test is recommended for individuals experiencing hormonal symptoms such as "
    "fatigue, mood swings, irregular cycles, or hot flashes. It is also useful for those "
    "with chronic stress or adrenal dysfunction, couples with fertility concerns, women "
    "in menopause or perimenopause, and athletes seeking performance optimization."
)
doc.add_paragraph(
    "The process involves four steps: (1) a consultation with a naturopathic doctor to "
    "assess symptoms and determine necessity, (2) collecting urine samples on filter "
    "paper strips at specified times throughout the day, (3) air drying the samples and "
    "shipping them to a specialized laboratory, and (4) laboratory analysis and "
    "interpretation of results with the patient. The DUTCH test can also be used to "
    "monitor the effectiveness of hormone therapy over time."
)

# ── 3. Micronutrient Testing ────────────────────────────────────────
doc.add_heading("Micronutrient Testing (MNT)", level=2)
doc.add_paragraph(
    "The Cellular Micronutrient Test is a cutting-edge diagnostic tool that measures and "
    "provides an in-depth analysis of 33 various vitamins and minerals in your body. "
    "Unlike standard serum-only tests, this test assesses intracellular nutrient "
    "functionality, taking into account factors such as age, genetics, individual "
    "biochemistry, environmental stressors, and the cumulative effects of diet over time."
)
doc.add_paragraph(
    "The test evaluates approximately 33 micronutrients including vitamins (C, D, "
    "B-vitamins), minerals (calcium, magnesium, zinc, iron), and additional nutrients "
    "supporting cellular function."
)
doc.add_paragraph(
    "This test is recommended for individuals with unexplained chronic symptoms, those "
    "on restrictive diets (vegan, vegetarian, food allergies), athletes and fitness "
    "enthusiasts, pregnant and nursing women, people with chronic health conditions, and "
    "anyone seeking optimal nutritional status."
)
doc.add_paragraph(
    "The process involves a consultation, maintaining your regular diet with fasting if "
    "instructed, a blood sample collection at the clinic, laboratory analysis using "
    "advanced techniques, and a personalized report comparing results to reference ranges."
)

# ── 4. SIBO Breath Test ─────────────────────────────────────────────
doc.add_heading("SIBO Breath Test (Lactulose Breath Test)", level=2)
doc.add_paragraph(
    "The Lactulose Breath Test, also called the SIBO Breath Test, is a non-invasive "
    "diagnostic procedure designed to detect Small Intestinal Bacterial Overgrowth (SIBO) "
    "by measuring hydrogen and methane gases in exhaled breath. These gases indicate the "
    "presence and type of bacterial overgrowth in the small intestine."
)
doc.add_paragraph(
    "This test is recommended for individuals experiencing persistent digestive symptoms "
    "including bloating, abdominal pain, diarrhea, or constipation, as well as those "
    "with an IBS diagnosis who may have underlying SIBO."
)
doc.add_paragraph(
    "Preparation involves following a special diet for 1–2 days prior and fasting "
    "overnight. A lactulose solution is consumed — it is not absorbed in the small "
    "intestine but is fermented by bacteria. Breath samples are then collected at regular "
    "intervals (typically every 20 minutes) for several hours using specialized devices."
)
doc.add_paragraph(
    "Results are classified as: hydrogen-dominant SIBO (elevated hydrogen levels), "
    "methane-dominant SIBO (elevated methane levels), or mixed SIBO (both gases elevated). "
    "The current clinic price for this test is $205. A naturopathic consultation is "
    "required before testing."
)

# ── 5. Comprehensive Gut Testing ────────────────────────────────────
doc.add_heading("Comprehensive Gut Testing (GI-360)", level=2)
doc.add_paragraph(
    "The GI-360 Test is one of the most comprehensive gut assessment tools available for "
    "evaluating multiple aspects of gastrointestinal health. It examines six key areas: "
    "digestive function (stomach acid, enzyme activity, nutrient absorption), microbiome "
    "balance, inflammation markers, immune function, parasites and pathogens, and food "
    "sensitivities."
)
doc.add_paragraph(
    "This test is recommended for individuals with digestive disorders or discomfort, "
    "those with suspected nutrient deficiencies, patients with autoimmune disease "
    "concerns, and anyone seeking a preventive gastrointestinal assessment."
)
doc.add_paragraph(
    "Sample types include stool samples (primary), with urine and blood samples sometimes "
    "required. Patients may need to follow specific dietary restrictions for several days "
    "before testing as directed by their naturopathic doctor. Results include a detailed "
    "report with personalized treatment recommendations including diet, lifestyle "
    "modifications, supplementation, botanical medicines, and IV therapy options."
)

# ── 6. Food Allergy Testing ─────────────────────────────────────────
doc.add_heading("Serum Food Allergy Test (IgE)", level=2)
doc.add_paragraph(
    "The Serum Food Allergy Test measures Immunoglobulin E (IgE) antibodies to identify "
    "which specific foods trigger allergic reactions. Unlike IgG food sensitivity testing, "
    "IgE testing detects true allergic responses. IgE antibodies are designed to protect "
    "your body from parasites and play a key role in immediate allergic reactions."
)
doc.add_paragraph(
    "Common allergens tested include peanuts, tree nuts (almonds, walnuts), milk, eggs, "
    "fish, shellfish, soy, and wheat. The clinic offers multiple panel options: Food "
    "Allergy Panel, Inhalant Allergy Panel, and Component Allergy Panel (which identifies "
    "specific proteins within allergens for more precise results)."
)
doc.add_paragraph(
    "A positive result (elevated IgE antibodies) indicates a food allergy. A negative "
    "result does not rule out non-IgE food sensitivities — for those, the IgG Food "
    "Sensitivity Test may be more appropriate. The test requires a blood sample drawn "
    "during a consultation with a naturopathic doctor."
)

# ── 7. Pathogen Testing ─────────────────────────────────────────────
doc.add_heading("Pathogen Testing (PAIRS / Cyrex Array 12)", level=2)
doc.add_paragraph(
    "The Pathogen Associated Immunity Reactivity Screen (PAIRS), also known as Cyrex "
    "Array 12, is a functional blood test that determines if you have been exposed to "
    "pathogens — particularly ones that cause chronic illness. It examines whether your "
    "immune system has recognized and reacted to stealth pathogens by analyzing antibody "
    "levels in your blood."
)
doc.add_paragraph(
    "This test measures immune response history to various chronic pathogens, antibody "
    "levels indicating past exposure, immune system reactivity to bacteria, viruses, and "
    "fungi, risk factors for autoimmune diseases, and allergies and sensitivities."
)
doc.add_paragraph(
    "It is recommended for patients suspected of having chronic infections, those with "
    "autoimmune disease concerns, compromised immunity, immune system disorders, "
    "unexplained chronic symptoms, or allergy management needs."
)
doc.add_paragraph(
    "The test requires a blood draw from the arm. Results are compared to known pathogen "
    "databases to identify which chronic pathogens your immune system has encountered and "
    "the strength of your response. Clinical applications include diagnosing autoimmune "
    "diseases, monitoring immune health over time, tailoring personalized treatment plans, "
    "identifying specific allergens, and assessing vaccine response effectiveness."
)

# ── 8. Cardiovascular Testing ───────────────────────────────────────
doc.add_heading("Cardiovascular Testing (Serum Cardiometabolic Profile)", level=2)
doc.add_paragraph(
    "The Serum Cardiometabolic Profile is a comprehensive blood test that evaluates heart "
    "and metabolic health through multiple markers, going beyond basic lipid testing to "
    "assess cardiovascular disease and metabolic disorder risk."
)
doc.add_paragraph(
    "Cardiovascular markers measured include LDL and VLDL cholesterol, non-HDL "
    "lipoprotein cholesterol, oxidized LDL, small dense LDL, lipoprotein(a), "
    "apolipoprotein AI and B, high-sensitivity C-reactive protein (hsCRP), and "
    "lipoprotein-associated phospholipase-A2 (PLAC). Metabolic markers include fasting "
    "glucose and insulin, 1,5-anhydroglucitol (Glycomark), adiponectin, leptin, "
    "leptin-to-adiponectin ratio (LAR), homocysteine, and fibrinogen."
)
doc.add_paragraph(
    "This test is recommended for individuals with a family history of heart disease, "
    "stroke, or diabetes, and those at risk for cardiovascular disease or metabolic "
    "disorders. Preparation requires an 8–12 hour fast (water permitted). The test "
    "requires a blood draw from the arm vein and provides a comprehensive analysis of "
    "15+ cardiovascular and metabolic biomarkers with clinical interpretation."
)

# ── 9. Autoimmunity Testing ─────────────────────────────────────────
doc.add_heading("Autoimmunity Testing (Antibody Array 5)", level=2)
doc.add_paragraph(
    "The Multiple Autoimmune Reactivity Screen, also known as Antibody Array 5 (AAb5), "
    "uses cutting-edge microarray technology to simultaneously assess autoimmune "
    "reactivity across multiple antigens in a single blood sample. Rather than ordering "
    "numerous individual tests, this multiplex approach screens for autoantibodies the "
    "immune system produces when attacking the body's own tissues."
)
doc.add_paragraph(
    "The test detects reactivity to diverse markers including Antinuclear Antibodies "
    "(ANAs), Anti-CCP, Anti-Smooth Muscle, Anti-Jo-1, Anti-tTG, Anti-dsDNA, Anti-AChR, "
    "Anti-Scl-70, Anti-GAD, and Anti-CCP2 antibodies. It can help diagnose rheumatoid "
    "arthritis, lupus (SLE), Sjogren's syndrome, systemic sclerosis, celiac disease, "
    "myasthenia gravis, type 1 diabetes, autoimmune hepatitis, and other conditions."
)
doc.add_paragraph(
    "This test is recommended for individuals with unexplained symptoms suggesting "
    "autoimmune disease, those experiencing diagnostic delays from traditional methods, "
    "patients seeking early detection of autoimmune reactivity, and those with known "
    "autoimmune conditions needing comprehensive assessment."
)
doc.add_paragraph(
    "Key advantages include early detection capability due to high sensitivity, reduced "
    "diagnostic delays compared to traditional methods requiring multiple tests, "
    "customizable antigen panels tailored to individual history, and cost-effectiveness "
    "as a single test versus multiple individual tests. Only one blood draw is required."
)

# ── 10. Gluten Testing ──────────────────────────────────────────────
doc.add_heading("Gluten Testing (Cyrex Array 3x)", level=2)
doc.add_paragraph(
    "The Wheat/Gluten Proteome Reactivity and Autoimmunity Test, known as Cyrex Array 3x, "
    "measures immune reactivity to wheat and gluten proteins. It helps identify various "
    "forms of gluten-related reactions beyond standard celiac disease testing."
)
doc.add_paragraph(
    "The test assesses reactivity to multiple wheat proteins and peptides including "
    "anti-tissue transglutaminase (tTG) antibodies, anti-endomysial (EMA) antibodies, "
    "deamidated gliadin peptide (DGP) antibodies, and wheat-specific IgE antibodies "
    "for allergy differentiation."
)
doc.add_paragraph(
    "This test is recommended for individuals with unexplained chronic gastrointestinal "
    "symptoms, first-degree relatives of celiac disease patients, those with autoimmune "
    "conditions (type 1 diabetes, thyroid disorders, rheumatoid arthritis), people "
    "experiencing fatigue, mood disturbances, or unexplained skin rashes, and those "
    "with suspected wheat allergies needing differentiation from other conditions."
)
doc.add_paragraph(
    "Testing involves a blood sample, medical evaluation, symptom assessment, and may "
    "include an elimination diet (gluten-free for weeks to months) followed by controlled "
    "reintroduction while monitoring symptoms."
)

# ── 11. Mold Testing ────────────────────────────────────────────────
doc.add_heading("Mold Testing (MycoTOX Profile)", level=2)
doc.add_paragraph(
    "The MycoTOX Profile is a specialized diagnostic urine test designed to detect and "
    "quantify mycotoxins (mold toxins) in the body. The test uses high-performance liquid "
    "chromatography (HPLC) and mass spectrometry (MS) to identify and quantify specific "
    "mycotoxin concentrations."
)
doc.add_paragraph(
    "This test is recommended for patients with mold allergies or sensitivities, those "
    "with suspected environmental mold exposure, individuals experiencing respiratory "
    "issues (coughing, wheezing), and patients with chronic illnesses such as "
    "fibromyalgia, chronic fatigue, or autoimmune disorders."
)
doc.add_paragraph(
    "Symptoms associated with mold exposure include respiratory symptoms, allergic "
    "reactions (sneezing, itchy eyes, skin rashes), fatigue and low energy, headaches "
    "and migraines, neurological symptoms (memory problems, concentration difficulties), "
    "skin issues, and immunosuppression."
)
doc.add_paragraph(
    "Treatment after positive results may include IV therapy with nutrients and "
    "antioxidants, supplementation (vitamin C, N-acetylcysteine, glutathione), dietary "
    "modifications, botanical medicines, and detoxification strategies. Identifying and "
    "remediating the environmental source of mold is equally important as treating "
    "symptoms. The test requires a urine sample collected during the consultation."
)

# ── 12. Environmental Toxin Testing ─────────────────────────────────
doc.add_heading("Environmental Toxin Testing (GPL-Tox)", level=2)
doc.add_paragraph(
    "The Environmental Toxicity Profile, known as GPL-Tox, is an advanced laboratory "
    "analysis examining markers in your body to identify toxin presence and assess "
    "detoxification capabilities. It evaluates exposure from environmental pollutants, "
    "food sources, and internal metabolic processes."
)
doc.add_paragraph(
    "The test measures specific toxin levels and exposure, detoxification pathway "
    "function, the body's capacity to eliminate toxins, and oxidative stress markers."
)
doc.add_paragraph(
    "Four profile options are available: (1) Basic — a broad overview of toxic exposure "
    "and detoxification markers, (2) Comprehensive — detailed analysis of heavy metals, "
    "chemicals, and organic pollutants with detoxification pathway evaluation, "
    "(3) Myco — focused on mycotoxins from mold exposure, and (4) Glyphosate — detects "
    "herbicide residues."
)
doc.add_paragraph(
    "This test is recommended for those with unexplained chronic health issues, "
    "individuals with suspected environmental toxin exposure, people considering "
    "detoxification programs, and anyone interested in preventive health measures. "
    "Sample types include urine and blood. Personalized treatment plans may include "
    "dietary modifications, supplements, or IV therapy to support detoxification pathways."
)

# ── Understanding Food Reactions ──────────────────────────────────────
doc.add_heading("Understanding Food Reactions — Allergies, Sensitivities, and Intolerances", level=2)
doc.add_paragraph(
    "Food reactions can be of three types: food allergies, food sensitivities, and "
    "food intolerances. While some people may think they have a food allergy, it might "
    "be a food intolerance instead, especially if the symptoms involve histamine. Since "
    "ALL adverse food reactions create inflammation in the body, it is important to "
    "identify which food reaction is occurring and treat it accordingly. Testing may or "
    "may not be required, but the naturopathic doctor will help identify the problem and "
    "may request a test if necessary."
)
doc.add_paragraph(
    "Food Allergies (IgE-mediated): These are immediate immune reactions triggered by "
    "IgE antibodies. Symptoms can appear within minutes and include hives, swelling, "
    "difficulty breathing, and in severe cases, anaphylaxis. Food allergies are commonly "
    "associated with histamine release — when the immune system overreacts, it triggers "
    "mast cells to release histamine, causing many of the classic allergy symptoms. "
    "Common allergens include peanuts, tree nuts, shellfish, milk, eggs, and wheat. "
    "The IgE Food Allergy Test at Nova Clinic detects these true allergic responses."
)
doc.add_paragraph(
    "Food Sensitivities (IgG-mediated): These are delayed immune reactions triggered by "
    "IgG antibodies. Unlike food allergies and intolerances, food sensitivities are not "
    "commonly associated with histamine. Symptoms may appear hours or even days after "
    "eating and include bloating, migraines, fatigue, joint pain, skin issues, and "
    "chronic inflammation. Because of the delay, food sensitivities are often harder to "
    "identify without testing. The IgG Food Sensitivity Test at Nova Clinic helps "
    "pinpoint these reactions."
)
doc.add_paragraph(
    "Food Intolerances (non-immune): These are reactions that do not involve the immune "
    "system directly. They are often caused by enzyme deficiencies (e.g. lactose "
    "intolerance), sensitivity to food chemicals (e.g. histamine, sulfites, caffeine), "
    "or other digestive issues. Symptoms typically involve the digestive system — gas, "
    "bloating, diarrhea, stomach pain — and can often be mistaken for food allergies, "
    "especially when histamine is involved. A naturopathic doctor can help distinguish "
    "intolerances from true allergies or sensitivities and recommend the right approach."
)

# ── Which Test Is Right For Me ───────────────────────────────────────
doc.add_heading("Functional Testing — Which Test Is Right For Me?", level=2)
doc.add_paragraph(
    "For digestive concerns: Consider the GI-360 Comprehensive Gut Test for a full "
    "gastrointestinal assessment, the SIBO Breath Test if bloating and IBS-like symptoms "
    "are the primary concern, or Food Sensitivity Testing if symptoms seem food-related."
)
doc.add_paragraph(
    "For hormonal concerns: The DUTCH Hormone Test provides the most comprehensive "
    "hormonal assessment available, measuring 18 markers across sex hormones, adrenal "
    "hormones, and neurotransmitters. For thyroid concerns specifically, blood testing "
    "is more appropriate — discuss with your naturopathic doctor."
)
doc.add_paragraph(
    "For unexplained chronic symptoms: Micronutrient Testing can reveal hidden "
    "deficiencies. The Autoimmunity Screen (Array 5) can detect early autoimmune "
    "reactivity. Mold Testing or Environmental Toxin Testing may be appropriate if "
    "environmental exposure is suspected."
)
doc.add_paragraph(
    "For cardiovascular risk: The Serum Cardiometabolic Profile provides 15+ markers "
    "beyond basic lipid panels, ideal for patients with family history of heart disease "
    "or metabolic conditions."
)
doc.add_paragraph(
    "For allergy and immune concerns: The IgE Food Allergy Test identifies true allergic "
    "reactions. Pathogen Testing (PAIRS/Array 12) assesses immune history to chronic "
    "pathogens. Gluten Testing (Array 3x) goes beyond standard celiac screening."
)
doc.add_paragraph(
    "All functional tests require an initial naturopathic consultation. If you're unsure "
    "which test is right for you, book a consultation and your practitioner will guide "
    "you to the appropriate testing based on your individual needs."
)

# ── Save ─────────────────────────────────────────────────────────────
out_path = Path("kb/sources/nova_clinic_functional_testing.docx")
doc.save(str(out_path))
print(f"Created {out_path} successfully!")

# Verify by parsing
from app.utils.docx_parser import parse_docx
sections = parse_docx(str(out_path))
print(f"Parsed {len(sections)} sections:")
for s in sections:
    print(f"  - {s['heading']} ({len(s['text'])} chars)")
