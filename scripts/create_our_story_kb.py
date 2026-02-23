#!/usr/bin/env python3
"""
Generate a .docx KB document from https://www.novaclinic.ca/our-story
"""
from docx import Document
from pathlib import Path

doc = Document()

# ── Clinic Identity & Philosophy ──────────────────────────────────────
doc.add_heading("ABOUT NOVA CLINIC — OUR STORY", level=1)

doc.add_heading("Who We Are", level=2)
doc.add_paragraph(
    "Nova Naturopathic Integrative Clinic is a Calgary South naturopathic clinic "
    "located at 208-6707 Elbow Drive Southwest, Calgary, AB T2V 0E4 "
    "(2nd floor, Mayfair Place building, commercial side). "
    "Nova Clinic won Calgary's Top Choice Alternative Health Services award in 2026."
)

doc.add_heading("Our Philosophy & Approach", level=2)
doc.add_paragraph(
    "Nova Clinic's foundational belief is that everyone has an inherent "
    "self-organizing, ordered healing process, which establishes, maintains "
    "and restores health. Our practitioners aim to support, facilitate, and "
    "augment this process by identifying and removing obstacles to health and "
    "recovery, and by supporting the creation of a healthy internal and external "
    "environment."
)
doc.add_paragraph(
    "The clinic integrates alternative and complementary strategies with "
    "conventional allopathic standards of care, rather than operating from a "
    "single therapeutic model. Treatment principles include: viewing patients "
    "holistically rather than symptom-focused; addressing root causes (the 'why' "
    "of a disease or illness rather than the 'what'); personalizing treatment "
    "plans through extensive case analysis; using gentle, minimally invasive "
    "modalities; and incorporating current scientific evidence."
)

# ── Professional Training Standards ───────────────────────────────────
doc.add_heading("Naturopathic Doctor Training Standards", level=2)
doc.add_paragraph(
    "Naturopathic doctors at Nova Clinic complete rigorous professional training: "
    "an undergraduate degree plus a four-year accredited doctorate program; "
    "a minimum of 4,100 classroom hours in medical sciences and therapeutics; "
    "a minimum of 1,200 supervised clinical practicum hours; "
    "provincial licensing board examinations; "
    "and 70 hours of continuing education per three-year cycle."
)

# ── Six Principles of Naturopathy ─────────────────────────────────────
doc.add_heading("The Six Principles of Naturopathy", level=2)
doc.add_paragraph(
    "Naturopathic medicine at Nova Clinic is guided by six core principles: "
    "1. Do No Harm — use the most natural, least invasive, and least toxic therapies. "
    "2. The Healing Power of Nature — trust the body's inherent ability to heal itself. "
    "3. Identify and Treat the Causes — look beyond symptoms to address root causes. "
    "4. Doctor as Teacher — educate patients and encourage self-responsibility for health. "
    "5. Treat the Whole Person — consider physical, mental, emotional, genetic, "
    "environmental, and social factors. "
    "6. Prevention — focus on overall health and disease prevention."
)

# ── What is Naturopathic Medicine ─────────────────────────────────────
doc.add_heading("What is Naturopathic Medicine?", level=2)
doc.add_paragraph(
    "Naturopathic medicine is a system of healthcare that utilizes natural methods "
    "and substances to restore and maintain health while preventing illness. "
    "Practitioners identify root causes and develop personalized treatment plans. "
    "They combine natural and conventional medical training as primary care "
    "physicians, with complete healing and wellness as the goal."
)
doc.add_paragraph(
    "Is naturopathic medicine safe? Yes. Multiple research studies demonstrate "
    "naturopathic treatments are both safe and effective for common conditions "
    "including heart disease, diabetes, chronic back pain, and anxiety."
)

doc.add_heading("What Happens at a Naturopath Appointment?", level=2)
doc.add_paragraph(
    "Initial visits typically include a detailed history of your current and past "
    "health concerns, a physical exam, and ordering laboratory tests to assess and "
    "diagnose your health status. Your naturopathic doctor will then develop a "
    "personalized treatment plan based on the findings."
)

# ── Save ──────────────────────────────────────────────────────────────
out_path = Path("kb/sources/nova_clinic_our_story.docx")
doc.save(str(out_path))
print(f"Created {out_path} successfully!")
